from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
from typing import List

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from robust_processing import AnalysisPackage, OUTLIER_METHOD_LABELS, controller_convention_text

ACCENT = RGBColor(0x0F, 0x4C, 0x81)
DARK = RGBColor(0x22, 0x22, 0x22)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_border(cell, color: str = "D8E1EB") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _style_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)


def _make_header(table, labels: List[str]) -> None:
    hdr = table.rows[0].cells
    for idx, label in enumerate(labels):
        hdr[idx].text = label
        _set_cell_shading(hdr[idx], "E9F1F8")
        for p in hdr[idx].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = DARK


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True


def _ranking_plot(entries: List[AnalysisPackage], path: Path) -> None:
    labels = [e.meta.get("loop_name", f"Malha {i+1}") for i, e in enumerate(entries)]
    scores = [e.loop_score for e in entries]
    order = np.argsort(scores)
    labels = [labels[i] for i in order]
    scores = [scores[i] for i in order]
    plt.figure(figsize=(8.4, max(3.0, 0.55 * len(labels) + 1.0)))
    y = np.arange(len(labels))
    plt.barh(y, scores)
    plt.yticks(y, labels)
    plt.xlim(0, 10)
    plt.xlabel("Nota da malha (0 a 10)")
    plt.title("Ranking consolidado do projeto")
    plt.grid(axis="x", alpha=0.25)
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight")
    plt.close()


def _plot_raw_clean(entry: AnalysisPackage, path: Path) -> None:
    raw = entry.raw_data
    clean = entry.clean_data
    plt.figure(figsize=(9.0, 4.3))
    plt.plot(raw["time"], raw["pv"], label="PV bruta", alpha=0.45)
    plt.plot(clean["time"], clean["pv"], label="PV limpa", linewidth=1.8)
    if "sp" in raw.columns:
        plt.plot(raw["time"], raw["sp"], label="SP", linewidth=1.2, linestyle="--")
    if "outlier_flag" in raw.columns:
        mask = raw["outlier_flag"] == 1
        if mask.any():
            plt.scatter(raw.loc[mask, "time"], raw.loc[mask, "pv"], label="Pontos sinalizados", s=10)
    plt.title("PV bruta x PV limpa")
    plt.grid(alpha=0.25)
    plt.legend(fontsize=8)
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight")
    plt.close()


def _plot_outlier_methods(entry: AnalysisPackage, path: Path) -> None:
    labels = [OUTLIER_METHOD_LABELS.get(m.method, m.method) for m in entry.outlier_methods]
    vals = [m.validation_fit_pct for m in entry.outlier_methods]
    plt.figure(figsize=(8.8, 4.1))
    plt.bar(range(len(labels)), vals)
    plt.xticks(range(len(labels)), labels, rotation=20, ha="right")
    plt.ylabel("FIT de validação (%)")
    plt.title("Comparação dos métodos de outlier")
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight")
    plt.close()


def _plot_model_validation(entry: AnalysisPackage, path: Path) -> None:
    dfv = entry.model_validation_series[entry.selected_model_name]
    plt.figure(figsize=(9.0, 4.2))
    plt.plot(dfv["time"], dfv["pv_real"], label="PV real", linewidth=1.8)
    plt.plot(dfv["time"], dfv["pv_pred"], label="PV prevista", linewidth=1.5)
    plt.title(f"PV real x PV prevista — {entry.selected_model_name}")
    plt.grid(alpha=0.25)
    plt.legend(fontsize=8)
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight")
    plt.close()


def _plot_tuning_pv(entry: AnalysisPackage, path: Path) -> None:
    baseline_pv = float(entry.clean_data["pv"].median()) if len(entry.clean_data) else 0.0
    baseline_sp = float(entry.clean_data["sp"].median()) if "sp" in entry.clean_data.columns else baseline_pv
    plt.figure(figsize=(9.1, 4.7))
    first = None
    for name, payload in entry.simulation_results.items():
        sim = payload["simulation"].data
        if first is None:
            first = sim
        plt.plot(sim["time"] / 60.0, sim["pv"] + baseline_pv, label=name)
    if first is not None:
        plt.plot(first["time"] / 60.0, first["sp"] + baseline_sp, "--", linewidth=1.4, label="SP previsto")
    plt.axhline(baseline_pv, linestyle=":")
    plt.xlabel("Tempo (min)")
    plt.ylabel("PV em unidades reais")
    plt.title("Como a PV deve se comportar para cada sintonia")
    plt.grid(alpha=0.25)
    plt.legend(fontsize=8)
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight")
    plt.close()


def _plot_tuning_mv(entry: AnalysisPackage, path: Path) -> None:
    plt.figure(figsize=(9.1, 4.2))
    for name, payload in entry.simulation_results.items():
        sim = payload["simulation"].data
        plt.plot(sim["time"] / 60.0, sim["mv"], label=name)
    plt.xlabel("Tempo (min)")
    plt.ylabel("MV")
    plt.title("Esforço da MV previsto para cada sintonia")
    plt.grid(alpha=0.25)
    plt.legend(fontsize=8)
    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches="tight")
    plt.close()


def build_report_docx(entries: List[AnalysisPackage], output_path: str | Path) -> Path:
    output_path = Path(output_path)
    doc = Document()
    _configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Relatório Técnico — Diagnóstico, Modelagem e Sintonia PID")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Versão de campo v8: 5 métodos de outlier, múltiplos modelos, PV real x PV prevista e comparação clara das sintonias.").italic = True

    with TemporaryDirectory() as td:
        td = Path(td)
        if len(entries) > 1:
            doc.add_heading("1. Ranking consolidado", level=1)
            rank_path = td / "ranking.png"
            _ranking_plot(entries, rank_path)
            doc.add_picture(str(rank_path), width=Inches(6.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, entry in enumerate(entries, start=1):
            if idx > 1 or len(entries) > 1:
                doc.add_page_break()

            conv = controller_convention_text(entry.equivalent_model.gain)
            doc.add_heading(f"{idx}. {entry.meta.get('loop_name', f'Malha {idx}')}", level=1)
            p = doc.add_paragraph()
            p.add_run("Cliente: ").bold = True
            p.add_run(str(entry.meta.get("client", "-")))
            p.add_run(" | Planta: ").bold = True
            p.add_run(str(entry.meta.get("site", "-")))
            p.add_run(" | Tag: ").bold = True
            p.add_run(str(entry.meta.get("loop_tag", "-")))
            p.add_run(" | Tipo: ").bold = True
            p.add_run(str(entry.meta.get("loop_type", "-")))

            p = doc.add_paragraph()
            p.add_run("Nota da malha: ").bold = True
            p.add_run(f"{entry.loop_score:.1f}/10 ({entry.grade})")
            p.add_run(" | Método de outlier selecionado: ").bold = True
            p.add_run(OUTLIER_METHOD_LABELS.get(entry.selected_outlier_method, entry.selected_outlier_method))
            p.add_run(" | Modelo-base: ").bold = True
            p.add_run(entry.selected_model_name)

            doc.add_paragraph(conv["note"])

            doc.add_heading("Resumo executivo", level=2)
            for reason in entry.reasons:
                doc.add_paragraph(reason, style="List Bullet")

            met = entry.performance_metrics
            summary_table = doc.add_table(rows=1, cols=4)
            _make_header(summary_table, ["Indicador", "Valor", "Indicador", "Valor"])
            rows = [
                ("RMSE SP-PV", f"{met['rmse']:.4f}", "Desvio padrão PV", f"{met['pv_std']:.4f}"),
                ("PV em ±0,05", f"{met['within_005_pct']:.1f}%", "PV em ±0,10", f"{met['within_010_pct']:.1f}%"),
                ("Mov. médio MV", f"{met['mv_move_mean']:.4f}", "Saturação MV", f"{0 if np.isnan(met['sat_frac_pct']) else met['sat_frac_pct']:.2f}%"),
                ("Ganho do processo (assinado)", f"{entry.equivalent_model.gain:.6f}", "Convenção do erro", conv["error_formula"]),
            ]
            for a, b, c, d in rows:
                row = summary_table.add_row().cells
                row[0].text, row[1].text, row[2].text, row[3].text = a, b, c, d
            _style_table(summary_table)

            raw_path = td / f"raw_clean_{idx}.png"
            outlier_path = td / f"outlier_{idx}.png"
            val_path = td / f"validation_{idx}.png"
            pv_path = td / f"pv_{idx}.png"
            mv_path = td / f"mv_{idx}.png"
            _plot_raw_clean(entry, raw_path)
            _plot_outlier_methods(entry, outlier_path)
            _plot_model_validation(entry, val_path)
            _plot_tuning_pv(entry, pv_path)
            _plot_tuning_mv(entry, mv_path)

            doc.add_heading("Qualidade do dado", level=2)
            doc.add_picture(str(raw_path), width=Inches(6.35))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(outlier_path), width=Inches(6.1))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            outlier_table = doc.add_table(rows=1, cols=4)
            _make_header(outlier_table, ["Método", "% sinalizado", "RMSE validação", "FIT validação (%)"])
            for m in entry.outlier_methods:
                row = outlier_table.add_row().cells
                row[0].text = OUTLIER_METHOD_LABELS.get(m.method, m.method)
                row[1].text = f"{m.flagged_pct:.2f}%"
                row[2].text = f"{m.validation_rmse:.5f}"
                row[3].text = f"{m.validation_fit_pct:.2f}"
            _style_table(outlier_table)

            doc.add_heading("Modelagem do processo", level=2)
            model_table = doc.add_table(rows=1, cols=4)
            _make_header(model_table, ["Modelo", "RMSE validação", "FIT validação (%)", "R² validação"])
            for m in entry.model_results:
                row = model_table.add_row().cells
                row[0].text = m.model_name
                row[1].text = f"{m.validation_rmse:.5f}"
                row[2].text = f"{m.validation_fit_pct:.2f}"
                row[3].text = f"{m.validation_r2:.3f}"
            _style_table(model_table)
            doc.add_picture(str(val_path), width=Inches(6.35))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(
                f"Modelo equivalente usado para a sintonia: ganho = {abs(entry.equivalent_model.gain):.6f}, "
                f"tau = {entry.equivalent_model.tau:.1f} s, tempo morto = {entry.equivalent_model.dead_time:.1f} s."
            )

            doc.add_heading("Sintonias propostas", level=2)
            tune_table = doc.add_table(rows=1, cols=9)
            _make_header(tune_table, [
                "Método", "Família", "Tipo", "Kc", "PB eq. (%)", "Ti (s)", "Ki (1/s)", "Td (s)", "TV(MV)"
            ])
            for _, row_df in entry.tuning_table.iterrows():
                sim_metrics = entry.simulation_results[row_df["name"]]["simulation"].metrics
                row = tune_table.add_row().cells
                row[0].text = str(row_df["name"])
                row[1].text = str(row_df["method"])
                row[2].text = str(row_df["controller_type"])
                row[3].text = f"{row_df['kc']:.4f}"
                row[4].text = f"{row_df['pb_percent_if_normalized']:.2f}"
                row[5].text = f"{row_df['ti_seconds']:.1f}"
                row[6].text = f"{row_df['ki_per_s']:.6f}"
                row[7].text = f"{row_df['td']:.1f}"
                row[8].text = f"{sim_metrics['total_variation_mv']:.1f}"
            _style_table(tune_table)
            doc.add_paragraph(
                "Kc é ganho do controlador. Ti e Td estão em segundos. Ki = Kc/Ti em 1/s. "
                "PB equivalente só vale se PV e MV estiverem normalizados em %."
            )
            doc.add_picture(str(pv_path), width=Inches(6.35))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(mv_path), width=Inches(6.35))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            metrics_table = doc.add_table(rows=1, cols=5)
            _make_header(metrics_table, ["Método", "Overshoot (%)", "Settling (s)", "IAE", "ITAE"])
            for name, payload in entry.simulation_results.items():
                met_sim = payload["simulation"].metrics
                row = metrics_table.add_row().cells
                row[0].text = name
                row[1].text = f"{met_sim['overshoot_pct']:.1f}"
                row[2].text = f"{met_sim['settling_time']:.0f}"
                row[3].text = f"{met_sim['IAE']:.3f}"
                row[4].text = f"{met_sim['ITAE']:.3f}"
            _style_table(metrics_table)

            doc.add_heading("Leitura prática para campo", level=2)
            doc.add_paragraph(
                "1. Use primeiro a curva da PV prevista para explicar ao cliente como a malha deve responder.\n"
                "2. Em seguida, mostre o gráfico da MV prevista para discutir esforço de válvula/atuador.\n"
                "3. Só depois entre nos números de Kc, Ti e Td.\n"
                "4. Se a planta tiver limites operacionais rígidos, comece pelo método mais robusto."
            )

            doc.add_heading("Checklist antes do teste em planta", level=2)
            for item in [
                "Confirmar que a malha está em AUTO durante o teste.",
                "Confirmar a escala real da MV e se existem travas, split-range ou limitações mecânicas.",
                "Checar se a convenção do erro no controlador é compatível com o sinal do ganho identificado.",
                "Verificar ruído do transmissor e filtragem já aplicada na PV.",
                "Executar primeiro um passo pequeno e observar segurança operacional antes de aumentar a agressividade.",
            ]:
                doc.add_paragraph(item, style="List Bullet")

            doc.add_heading("Recomendação objetiva para esta malha", level=2)
            first_name = list(entry.simulation_results.keys())[0] if entry.simulation_results else "—"
            doc.add_paragraph(
                f"Para esta execução, a recomendação inicial é testar primeiro a sintonia '{first_name}', porque ela costuma combinar "
                "menor risco operacional com uma curva de PV mais explicável para o cliente. "
                "Depois disso, compare em campo com a segunda melhor opção e valide se o esforço da MV permanece aceitável."
            )

    doc.save(output_path)
    return output_path
